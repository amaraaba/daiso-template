#!/usr/bin/env python3
"""
Word Document (.docx) to Markdown Importer

Converts .docx files into Markdown with YAML front-matter.
Saves to packs/<domain>/knowledge/guides/

Usage:
    py common/tools/docx_to_markdown_import.py --docx <path> --domain <name>

Example:
    py common/tools/docx_to_markdown_import.py --docx "path/to/document.docx" --domain nvl-test
"""

import argparse
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional
import re

try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def sanitize_filename(name: str) -> str:
    """Remove invalid filename characters."""
    return re.sub(r'[<>:"/\\|?*]', "_", name).strip("._")


def extract_docx_text(docx_path: Path) -> str:
    """
    Extract text from .docx file preserving structure.
    Returns: Markdown-formatted text
    """
    doc = Document(docx_path)
    lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        
        # Detect heading levels from style
        style_name = para.style.name if para.style else ""
        
        if "Heading 1" in style_name:
            lines.append(f"# {text}")
        elif "Heading 2" in style_name:
            lines.append(f"## {text}")
        elif "Heading 3" in style_name:
            lines.append(f"### {text}")
        else:
            lines.append(text)
    
    # Handle tables if present
    for table in doc.tables:
        lines.append("")
        lines.append("| " + " | ".join([cell.text.strip() for cell in table.rows[0].cells]) + " |")
        lines.append("|" + "|".join(["---|"] * len(table.rows[0].cells)))
        
        for row in table.rows[1:]:
            lines.append("| " + " | ".join([cell.text.strip() for cell in row.cells]) + " |")
        lines.append("")
    
    return "\n".join(lines)


def docx_to_markdown(docx_path: Path, section_name: str) -> tuple[str, str]:
    """
    Convert .docx to Markdown.
    Returns: (title, markdown_content)
    """
    # Extract title from filename
    title = sanitize_filename(docx_path.stem)
    
    # Extract content
    markdown = extract_docx_text(docx_path)
    
    return title, markdown


def main():
    parser = argparse.ArgumentParser(
        description="Convert Word documents (.docx) to Markdown files"
    )
    parser.add_argument(
        "--docx",
        required=True,
        type=Path,
        help="Path to .docx file"
    )
    parser.add_argument(
        "--domain",
        required=True,
        help="Domain/pack name (e.g. nvl-test)"
    )
    parser.add_argument(
        "--section-name",
        default=None,
        help="Section name to use in front-matter (default: document name)"
    )
    parser.add_argument(
        "--pack-root",
        default=".",
        type=Path,
        help="Pack root directory (default: current dir)"
    )
    
    args = parser.parse_args()
    docx_path = args.docx.resolve()
    pack_root = args.pack_root.resolve()
    
    # Validate
    if not docx_path.exists():
        print(f"ERROR: File not found: {docx_path}")
        return 1
    
    if not (pack_root / "packs").exists():
        print(f"ERROR: packs/ not found in {pack_root}")
        print("Run from your DAISO repository root.")
        return 1
    
    print(f"\n📄 Reading: {docx_path.name}")
    
    # Extract section name if not provided
    section_name = args.section_name or docx_path.stem
    
    # Convert
    try:
        title, markdown = docx_to_markdown(docx_path, section_name)
    except Exception as e:
        print(f"ERROR reading document: {e}")
        return 1
    
    # Create output directory
    output_dir = pack_root / "packs" / args.domain / "knowledge" / "guides" / sanitize_filename(section_name)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"💾 Output: {output_dir}")
    
    # Create YAML front-matter
    now = datetime.utcnow().isoformat()
    front_matter = f"""---
title: {title}
source: Word Document
section: {section_name}
imported: {now}
---

"""
    
    # Write file
    output_path = output_dir / f"{sanitize_filename(title)}.md"
    output_path.write_text(front_matter + markdown, encoding="utf-8")
    
    print(f"\n✅ Converted successfully!")
    print(f"\n📄 File saved:")
    rel_path = output_path.relative_to(pack_root)
    print(f"   {rel_path}")
    
    print(f"\n💡 Next steps:")
    print(f"   1. Review the Markdown file above")
    print(f"   2. git add packs/{args.domain}/knowledge/")
    print(f"   3. git commit -m \"Add {section_name}\"")
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
